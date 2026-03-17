from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY     = RGBColor(0x1F, 0x4E, 0x79)
MED_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY     = RGBColor(0x60, 0x60, 0x60)

doc = Document()

sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.left_margin   = sec.right_margin  = Inches(0.75)
sec.top_margin    = Inches(0.6)
sec.bottom_margin = Inches(0.45)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(8.5)

# ── Helpers ──────────────────────────────────────────────────────────

def sp(para, bef=0, aft=0, ln=None):
    pf = para.paragraph_format
    pf.space_before = Pt(bef)
    pf.space_after  = Pt(aft)
    if ln:
        pf.line_spacing = Pt(ln)

def run(para, text, bold=False, italic=False, sz=None, color=None):
    r = para.add_run(text)
    r.font.name = "Calibri"
    if bold:   r.font.bold   = True
    if italic: r.font.italic = True
    if sz:     r.font.size   = Pt(sz)
    if color:  r.font.color.rgb = color
    return r

def bottom_border(para, color="2E75B6", sz=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def no_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def cell_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def cell_margins(cell, top=0, left=0, bottom=0, right=0):
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)

def hdr(cell, text):
    p = cell.add_paragraph()
    sp(p, bef=6, aft=1, ln=10)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.bold = True
    r.font.size = Pt(7.5); r.font.color.rgb = MED_BLUE
    r.font.all_caps = True
    bottom_border(p, color="2E75B6", sz=4)

def body(cell, text, bold=False):
    p = cell.add_paragraph()
    sp(p, bef=1, aft=1, ln=9.5)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(7.5)
    if bold: r.font.bold = True

def bullet(cell, text, prefix=None):
    p = cell.add_paragraph()
    sp(p, bef=0, aft=0, ln=9.5)
    p.paragraph_format.left_indent       = Inches(0.14)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    dot = p.add_run("\u2022  ")
    dot.font.name = "Calibri"; dot.font.size = Pt(7.5); dot.font.color.rgb = MED_BLUE
    if prefix:
        rb = p.add_run(prefix + " ")
        rb.font.name = "Calibri"; rb.font.size = Pt(7.5); rb.font.bold = True
    rt = p.add_run(text)
    rt.font.name = "Calibri"; rt.font.size = Pt(7.5)

def numbered(cell, n, text, prefix=None):
    p = cell.add_paragraph()
    sp(p, bef=1, aft=0, ln=9.5)
    p.paragraph_format.left_indent       = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    rn = p.add_run(f"{n}. ")
    rn.font.name = "Calibri"; rn.font.size = Pt(7.5)
    rn.font.bold = True; rn.font.color.rgb = MED_BLUE
    if prefix:
        rb = p.add_run(prefix + " ")
        rb.font.name = "Calibri"; rb.font.size = Pt(7.5); rb.font.bold = True
    rt = p.add_run(text)
    rt.font.name = "Calibri"; rt.font.size = Pt(7.5)

# ═══════════════════════════════════════════════════════════════════
#  TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════

tp = doc.add_paragraph()
sp(tp, bef=0, aft=0, ln=24)
run(tp, "OAP Contest Manager", bold=True, sz=20, color=NAVY)
run(tp, "   \u2014   User Guide", sz=15, color=MED_BLUE)

st = doc.add_paragraph()
sp(st, bef=0, aft=5, ln=11)
run(st, "A Quick-Start Reference for UIL One-Act Play Contest Managers",
    italic=True, sz=9, color=GRAY)
bottom_border(st, color="1F4E79", sz=10)

# ═══════════════════════════════════════════════════════════════════
#  MAIN TWO-COLUMN TABLE
# ═══════════════════════════════════════════════════════════════════

tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
no_borders(tbl)

L = tbl.rows[0].cells[0]
R = tbl.rows[0].cells[1]
cell_width(L, 3.05)
cell_width(R, 3.95)
L.vertical_alignment = WD_ALIGN_VERTICAL.TOP
R.vertical_alignment = WD_ALIGN_VERTICAL.TOP
cell_margins(L, right=110)
cell_margins(R, left=110)
clear_cell(L)
clear_cell(R)

# ── LEFT — What It Is ───────────────────────────────────────────────
hdr(L, "What It Is")
body(L,
    "The OAP Contest Manager is a single HTML file that opens directly in your web "
    "browser\u2014no installation, no login, no internet connection required after the "
    "page loads. It collects all contest data in one place and generates a complete, "
    "formatted document package in seconds. Works best in Chrome or Edge.")

# ── LEFT — What It Produces ─────────────────────────────────────────
hdr(L, "What It Produces")
body(L, "Click Generate to download a ZIP file containing:")

docs_list = [
    "Year-Round Checklist",
    "Fall District Meeting Agenda",
    "Host School Checklist",
    "Rehearsal & Contest Schedule (Excel)",
    "Contest Day Schedule (Excel)",
    "Director Information Letter",
    ("Pre-Rehearsal Company Meeting",
     "\u2014 one-page checklist used during each company\u2019s rehearsal slot: stage manager, lights/sound, curtains, spike tape, show start, rules"),
    "Directors Meeting Script",
    "Awards Script",
    "Advancing Schools Letter",
    "School-Director Contact List",
    "Adjudicator Info Sheet",
    ("Adjudicator Packets (PDF)",
     "\u2014 official UIL evaluation, ranking & awards ballots, pre-filled per judge, merged into one print-ready file"),
    "Timer Instructions + Form",
]
for d in docs_list:
    if isinstance(d, tuple):
        bullet(L, d[1], prefix=d[0])
    else:
        bullet(L, d)

# ── LEFT — Tips ─────────────────────────────────────────────────────
hdr(L, "Tips & Notes")
tips = [
    ("Share it:",
     "Copy the .html file to any computer\u2014it is fully self-contained."),
    ("Year to year:",
     "Snapshots remember school names and CM info but skip dates and judge details, so annual updates are fast."),
    ("Adj. Packets:",
     "Unchecked by default. Enable only once judges are finalized (takes a few extra seconds to build)."),
    ("Dark mode:",
     "Use the \U0001f319 Dark button (bottom-right corner) when working in a dark theatre. Preference is saved."),
]
for pfx, txt in tips:
    bullet(L, txt, prefix=pfx)

# ── RIGHT — How to Use It ───────────────────────────────────────────
hdr(R, "How to Use It")

steps = [
    ("Open the file:",
     "Double-click OAP Contest Setup.html. It opens in your browser with no login needed."),
    ("Jump bar:",
     "A row of shortcut buttons sits at the top of the page. Click any button to jump directly to that section."),
    ("CM Info:",
     "Verify your name, email, phone, address, and website. Pre-filled with Allen's details\u2014update as needed."),
    ("Contest Identity:",
     "Fill upon appointment: year, classification (1A\u20136A), level (District, Bi-District, etc.), district number, host school, venue name and address."),
    ("Contest Details:",
     "Fill after your planning meeting: contest date, Directors Meeting time, first show time, critique format, number of judges and schools, rehearsal dates, entry fee, and deadlines."),
    ("Adjudicators:",
     "Fill after contracting: judge names, phone numbers, and hotel details if applicable."),
    ("Schools & Directors:",
     "One row per school. Enter school name, director name(s), and email address(es). Multiple directors per school are supported."),
    ("Play Titles:",
     "Assign each school a performance slot number and enter the play title."),
    ("Snapshots:",
     'Save your data under a name (e.g., "2026 District 20") to reload next season. Export/import to move snapshots between computers.'),
    ("Choose your docs:",
     "In Documents to Generate, check only what you need. Use Check All / Uncheck All for speed."),
    ("Generate:",
     "Click Generate. The tool validates your data and downloads a named ZIP of all selected documents."),
    ("Email drafts:",
     "Expand Email Draft Composer for pre-written templates (announcement, deadline reminder, day-before, judge instructions). Copy subject and body, then paste into your email client."),
    ("Schedule preview:",
     "Expand Contest Day Schedule Preview for a live look at your contest timeline as data is filled in."),
    ("Critique randomizer:",
     "Expand Critique Assignment Randomizer to randomly assign schools to judges, respecting the UIL rule that the last-performing school cannot go to Judge 1."),
]
for i, (pfx, txt) in enumerate(steps, 1):
    numbered(R, i, txt, prefix=pfx)

# ═══════════════════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════════════════

fp = doc.sections[0].footer.paragraphs[0]
fp.clear()
sp(fp, bef=0, aft=0, ln=9)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run(
    "OAP Contest Manager  \u2502  For internal use  \u2502  Updated March 17, 2026"
)
rf.font.name = "Calibri"; rf.font.size = Pt(6.5)
rf.font.color.rgb = GRAY; rf.font.italic = True

# ── Save ────────────────────────────────────────────────────────────
out = r"C:\Users\Allen\Desktop\Claude CoWork\OAP Documents\output\OAP Contest Manager - User Guide.docx"
doc.save(out)
print("Saved:", out)
